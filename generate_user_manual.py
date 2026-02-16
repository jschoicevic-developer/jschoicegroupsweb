"""
JS Choice Group - Professional User Manual & Final Report Generator
Generates a comprehensive .docx document with professional formatting
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ─── CONSTANTS ──────────────────────────────────────────────
PRIMARY = RGBColor(0xAB, 0xB3, 0xF1)        # Lavender Blue
SECONDARY = RGBColor(0xF1, 0xAB, 0xAB)      # Coral Pink
DARK = RGBColor(0x2D, 0x37, 0x48)           # Dark Charcoal
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x1A, 0x20, 0x2C)
GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_BG = RGBColor(0xF7, 0xFA, 0xFC)
SUCCESS = RGBColor(0x68, 0xD3, 0x91)
ACCENT = RGBColor(0x5A, 0x67, 0xD8)

# HEX versions for table shading
PRIMARY_HEX = "ABB3F1"
SECONDARY_HEX = "F1ABAB"
DARK_HEX = "2D3748"
LIGHT_HEX = "F0F2FF"
WHITE_HEX = "FFFFFF"
ACCENT_HEX = "E8EAFF"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", 4)}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_styled_paragraph(doc, text, font_name="Calibri", font_size=11, bold=False,
                          color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          space_before=0, space_after=6, italic=False, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(24)
            run.font.color.rgb = DARK
        elif level == 2:
            run.font.size = Pt(18)
            run.font.color.rgb = DARK
        elif level == 3:
            run.font.size = Pt(14)
            run.font.color.rgb = ACCENT
    return h


def add_divider(doc, color_hex=PRIMARY_HEX):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, DARK_HEX)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.color.rgb = BLACK
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                set_cell_shading(cell, ACCENT_HEX)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table


def add_bullet_list(doc, items, font_size=10, color=BLACK):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(font_size)
        run.font.color.rgb = color
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15


def add_info_box(doc, title, content, color_hex=LIGHT_HEX):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, color_hex)

    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = DARK

    p2 = cell.add_paragraph()
    run2 = p2.add_run(content)
    run2.font.name = "Calibri"
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

    for edge in ["top", "bottom", "start", "end"]:
        set_cell_border(cell, **{edge: {"val": "single", "sz": "4", "color": PRIMARY_HEX}})

    doc.add_paragraph()


# ─── MAIN DOCUMENT GENERATOR ──────────────────────────────
def generate_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    add_styled_paragraph(doc, "JS CHOICE GROUP PTY LTD",
                          font_size=32, bold=True, color=DARK,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    add_styled_paragraph(doc, "Trading as JS Choice \u2013 Care and Support",
                          font_size=14, color=ACCENT, italic=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_divider(doc, PRIMARY_HEX)

    add_styled_paragraph(doc, "USER MANUAL & FINAL PROJECT REPORT",
                          font_size=22, bold=True, color=PRIMARY,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=16, space_after=8)

    add_styled_paragraph(doc, "Comprehensive Platform Documentation",
                          font_size=14, color=GRAY,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_divider(doc, SECONDARY_HEX)

    # Cover details table
    cover_data = [
        ["Document Type:", "User Manual & Final Report"],
        ["Version:", "1.2.2"],
        ["Date:", "February 16, 2026"],
        ["Prepared For:", "JS Choice Group Pty Ltd"],
        ["Platform:", "jschoicegroup.com.au"],
        ["Classification:", "Internal / Client Documentation"],
    ]
    table = doc.add_table(rows=len(cover_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(cover_data):
        c1 = table.rows[i].cells[0]
        c2 = table.rows[i].cells[1]
        c1.text = ""
        c2.text = ""
        r1 = c1.paragraphs[0].add_run(label)
        r1.font.name = "Calibri"
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = GRAY
        r2 = c2.paragraphs[0].add_run(value)
        r2.font.name = "Calibri"
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        c1.width = Inches(2.5)
        c2.width = Inches(3.5)

    for _ in range(3):
        doc.add_paragraph()

    add_styled_paragraph(doc, "CONFIDENTIAL",
                          font_size=10, bold=True, color=SECONDARY,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, "This document contains proprietary information about the JS Choice Group digital platform.",
                          font_size=9, color=GRAY,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Table of Contents", level=1)
    add_divider(doc)

    toc_items = [
        "1. Executive Summary",
        "2. Platform Overview",
        "3. Technology Stack",
        "4. Brand & Color System",
        "5. Frontend Pages Directory",
        "    5.1 Public Website Pages",
        "    5.2 Service Pages",
        "    5.3 Location Pages",
        "    5.4 NDIS Tools",
        "    5.5 Content Pages",
        "6. Admin Panel (CRM) \u2013 User Manual",
        "    6.1 Dashboard Overview",
        "    6.2 Lead Management",
        "    6.3 Referral Management",
        "    6.4 Blog Management",
        "    6.5 Gallery Management",
        "    6.6 Analytics & Reports",
        "    6.7 Settings & Profile",
        "7. Backend Modules & API",
        "    7.1 Authentication Module",
        "    7.2 Lead/CRM Module",
        "    7.3 Blog Module",
        "    7.4 Gallery Module",
        "    7.5 NDIS Tools Module",
        "    7.6 Email Module",
        "    7.7 Analytics Module",
        "8. Database Architecture",
        "9. Features Summary",
        "10. Security & Compliance",
        "11. Deployment & Infrastructure",
        "12. Support & Contact",
    ]
    for item in toc_items:
        indent = item.startswith("    ")
        text = item.strip()
        p = add_styled_paragraph(doc, text, font_size=10 if indent else 11,
                                  color=GRAY if indent else DARK,
                                  bold=not indent, space_after=3)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Executive Summary", level=1)
    add_divider(doc)

    add_styled_paragraph(
        doc,
        "JS Choice Group Pty Ltd (trading as JS Choice \u2013 Care and Support) is a registered NDIS provider "
        "based in Point Cook, Melbourne. Founded by Jan Fardowsi, the organization delivers ethical, culturally "
        "responsive, and participant-centred disability support services.",
        font_size=11, color=BLACK, space_after=8
    )

    add_styled_paragraph(
        doc,
        "This document serves as a comprehensive User Manual and Final Project Report for the JS Choice Group "
        "digital platform. It covers all frontend pages, backend modules, admin panel operations, database "
        "architecture, and deployment infrastructure. This guide is intended for administrators, team members, "
        "and stakeholders who need to understand, operate, or maintain the platform.",
        font_size=11, color=BLACK, space_after=12
    )

    add_info_box(doc,
                 "Platform At A Glance",
                 "51+ Website Pages  |  7 Admin Modules  |  38+ API Endpoints  |  10 Database Tables\n"
                 "Full CRM System  |  NDIS Tools Suite  |  Blog & Gallery  |  Email Automation")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 2. PLATFORM OVERVIEW
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. Platform Overview", level=1)
    add_divider(doc)

    add_styled_paragraph(
        doc,
        "The JS Choice Group platform is a modern, full-stack web application that serves two primary audiences:",
        font_size=11, space_after=8
    )

    add_styled_table(doc,
        ["Audience", "Interface", "Purpose"],
        [
            ["NDIS Participants & Families", "Public Website", "Learn about services, use NDIS tools, submit referrals, read blog content"],
            ["JS Choice Team / Admin", "Admin Panel (CRM)", "Manage leads, referrals, blog posts, gallery, analytics, and communications"],
        ],
        col_widths=[2.0, 1.5, 3.0]
    )

    add_heading_styled(doc, "Platform Architecture", level=2)
    add_styled_paragraph(
        doc,
        "The platform follows a modern monolithic architecture with clear separation of concerns:",
        font_size=11, space_after=8
    )

    add_bullet_list(doc, [
        "Frontend: Next.js 16 (React 19) with Server-Side Rendering (SSR) and Static Site Generation (SSG)",
        "Backend: Next.js API Routes (serverless functions)",
        "Database: Supabase (PostgreSQL) with Row-Level Security (RLS)",
        "Storage: Supabase Storage for image uploads",
        "Authentication: Supabase Auth with session management",
        "Email: Resend API for transactional emails",
        "Hosting: Vercel (auto-scaling, edge network)",
        "Domain: jschoicegroup.com.au",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 3. TECHNOLOGY STACK
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Technology Stack", level=1)
    add_divider(doc)

    add_styled_table(doc,
        ["Category", "Technology", "Version", "Purpose"],
        [
            ["Framework", "Next.js", "16.1.6", "Full-stack React framework with SSR/SSG"],
            ["UI Library", "React", "19.2.3", "Component-based user interface"],
            ["Language", "TypeScript", "5.x", "Type-safe JavaScript"],
            ["Styling", "Tailwind CSS", "4.x", "Utility-first CSS framework"],
            ["Animations", "Framer Motion", "12.31.0", "Smooth page & element animations"],
            ["Icons", "Lucide React", "0.563.0", "Modern icon library (500+ icons)"],
            ["Database", "Supabase (PostgreSQL)", "2.95.3", "Managed database with auth & storage"],
            ["Email", "Resend", "6.9.1", "Transactional email delivery"],
            ["Rich Editor", "TipTap", "3.19.0", "WYSIWYG editor for blog posts"],
            ["Charts", "Recharts", "3.7.0", "Data visualization for analytics"],
            ["PDF Export", "jsPDF", "4.1.0", "PDF generation for budget calculator"],
            ["Excel Export", "xlsx", "0.18.5", "Spreadsheet exports"],
            ["Image Processing", "Sharp", "0.34.5", "Server-side image optimization"],
            ["UI Components", "Radix UI", "1.4.3", "Accessible component primitives"],
            ["CSV Parsing", "PapaParse", "5.5.3", "CSV data processing"],
        ],
        col_widths=[1.2, 1.8, 0.8, 2.7]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 4. BRAND & COLOR SYSTEM
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Brand & Color System", level=1)
    add_divider(doc)

    add_styled_paragraph(
        doc,
        "The JS Choice Group platform uses a carefully designed color system that conveys trust, "
        "compassion, and professionalism. Below is the complete brand palette:",
        font_size=11, space_after=12
    )

    # Color swatches table
    color_table = doc.add_table(rows=10, cols=4)
    color_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    color_table.style = "Table Grid"

    color_headers = ["Color Name", "Hex Code", "Preview", "Usage"]
    for i, h in enumerate(color_headers):
        cell = color_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, DARK_HEX)

    colors_data = [
        ["Primary (Lavender Blue)", "#ABB3F1", PRIMARY_HEX, "Buttons, links, primary accents, CTAs"],
        ["Secondary (Coral Pink)", "#F1ABAB", SECONDARY_HEX, "Secondary buttons, highlights, accents"],
        ["Dark Charcoal", "#2D3748", DARK_HEX, "Headings, dark backgrounds, footer"],
        ["Foreground", "#1A202C", "1A202C", "Body text, labels"],
        ["Background", "#F7FAFC", "F7FAFC", "Page backgrounds"],
        ["Accent", "#E8EAFF", ACCENT_HEX, "Light card backgrounds, hover states"],
        ["Success (Green)", "#68D391", "68D391", "Success states, positive indicators"],
        ["Destructive (Red)", "#E53E3E", "E53E3E", "Error states, delete actions"],
        ["Muted", "#EDF2F7", "EDF2F7", "Muted backgrounds, disabled states"],
    ]

    for r_idx, (name, hex_code, swatch_hex, usage) in enumerate(colors_data):
        row = color_table.rows[r_idx + 1]
        # Name
        row.cells[0].text = ""
        r = row.cells[0].paragraphs[0].add_run(name)
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.bold = True
        # Hex
        row.cells[1].text = ""
        r = row.cells[1].paragraphs[0].add_run(hex_code)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Swatch
        row.cells[2].text = ""
        row.cells[2].paragraphs[0].add_run("          ")
        set_cell_shading(row.cells[2], swatch_hex)
        # Usage
        row.cells[3].text = ""
        r = row.cells[3].paragraphs[0].add_run(usage)
        r.font.name = "Calibri"
        r.font.size = Pt(9)

    doc.add_paragraph()

    add_heading_styled(doc, "Typography", level=2)
    add_styled_table(doc,
        ["Element", "Font Family", "Weight", "Usage"],
        [
            ["Headings (H1\u2013H6)", "Dosis", "200\u2013800 (Black)", "All section headings, titles, labels"],
            ["Body Text", "Poppins", "100\u2013900 (Regular)", "Paragraphs, descriptions, form text"],
            ["Code / Mono", "Geist Mono", "Regular", "Technical text, NDIS codes"],
        ],
        col_widths=[1.5, 1.5, 1.5, 2.0]
    )

    add_heading_styled(doc, "Design Principles", level=2)
    add_bullet_list(doc, [
        "Rounded corners throughout (2rem \u2013 2.5rem border radius for cards, full-round for buttons)",
        "Glass-morphism effects (backdrop-blur, transparency) for modern depth",
        "Framer Motion animations on scroll (fade-in, slide, scale transitions)",
        "Consistent spacing system (generous padding: 6\u20138\u201310 units, section gaps: 20\u201332 units)",
        "Fully responsive design across mobile, tablet, and desktop breakpoints",
        "Dark mode support available",
        "Buttons always use dark text (#1A202C) on primary/secondary backgrounds",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 5. FRONTEND PAGES DIRECTORY
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Frontend Pages Directory", level=1)
    add_divider(doc)

    add_styled_paragraph(doc, "The platform consists of 51+ public-facing pages organized into logical categories.",
                          font_size=11, space_after=10)

    # 5.1 Main Pages
    add_heading_styled(doc, "5.1 Public Website Pages", level=2)
    add_styled_table(doc,
        ["Page", "Route", "Description"],
        [
            ["Home", "/", "Landing page with hero slider, about, services, FAQ, areas served, CTA sections"],
            ["About Us", "/about-us", "Company story, founder profile, values, commitment, why choose us"],
            ["Contact Us", "/contact-us", "Contact form with fields for name, email, phone, message + map"],
            ["Referral", "/referral", "Multi-step referral form (6 sections) for NDIS participants"],
            ["Blog", "/blog", "Blog listing with search, categories, featured posts"],
            ["Blog Post", "/blog/[slug]", "Individual blog post with rich content, author, tags"],
            ["Gallery", "/gallery", "Photo gallery with multi-image cards and categories"],
            ["Career", "/career", "Careers/job opportunities page"],
            ["Resources", "/resources", "NDIS resources and helpful links"],
            ["Thank You", "/thank-you", "Confirmation page after form submission"],
        ],
        col_widths=[1.2, 1.5, 3.8]
    )

    # 5.2 Service Pages
    add_heading_styled(doc, "5.2 NDIS Service Pages (15 Pages)", level=2)
    add_styled_table(doc,
        ["#", "Service Page", "Route"],
        [
            ["1", "Assistance with Daily Life", "/assistance-with-daily-life"],
            ["2", "Psychosocial Recovery Coach", "/psychosocial-recovery-coach"],
            ["3", "Assistance with Nursing Care", "/assistance-with-nursing-care"],
            ["4", "Emergency Respite", "/emergency-respite"],
            ["5", "Group Centre Activities", "/group-centre-activities"],
            ["6", "Transportation Assistance", "/transportation-assistance"],
            ["7", "Community Participation", "/access-to-community-activities"],
            ["8", "Support Coordination", "/support-coordination"],
            ["9", "Client & Family Advocacy", "/client-and-family-advocacy"],
            ["10", "NDIS-Only Advocacy", "/client-and-family-advocacy-for-ndis-participants-only"],
            ["11", "NDIS Access Requests", "/ndis-access-requests"],
            ["12", "Innovative Community", "/innovative-community-participation-..."],
            ["13", "Employment & Education", "/employment-education"],
            ["14", "NDIS Accommodation", "/ndis-accommodation"],
            ["15", "NDIS Accommodation Geelong", "/ndis-accommodation-geelong"],
        ],
        col_widths=[0.3, 2.5, 3.7]
    )

    # 5.3 Location Pages
    add_heading_styled(doc, "5.3 Location-Based SEO Pages (16 Pages)", level=2)
    add_styled_paragraph(doc, "Each location page is optimized for local SEO, targeting NDIS participants in specific Melbourne suburbs.",
                          font_size=10, color=GRAY, space_after=8)

    locations = [
        "Point Cook", "Tarneit", "Shepparton", "Werribee",
        "Hoppers Crossing", "Truganina", "Craigieburn", "Williams Landing",
        "Laverton", "Altona", "Altona North", "Altona Meadows",
        "South Morang", "Footscray", "Lara", "Epping"
    ]
    loc_rows = []
    for i, loc in enumerate(locations):
        slug = loc.lower().replace(" ", "-")
        loc_rows.append([str(i+1), loc, f"/ndis-providers-{slug}"])

    add_styled_table(doc, ["#", "Suburb", "Route"], loc_rows, col_widths=[0.3, 2.0, 4.2])

    # 5.4 NDIS Tools
    add_heading_styled(doc, "5.4 NDIS Tools Suite (4 Pages)", level=2)
    add_styled_table(doc,
        ["Tool", "Route", "Description"],
        [
            ["Tools Landing", "/tools", "Overview of all available NDIS tools"],
            ["NDIS Price Guide", "/tools/ndis-price-guide", "Search 5,000+ NDIS support items with region-specific pricing"],
            ["Price Item Detail", "/tools/ndis-price-guide/[code]", "Detailed view of individual NDIS support item"],
            ["Budget Calculator", "/tools/ndis-budget-calculator", "Plan NDIS budget with add-to-cart, frequency, PDF export"],
            ["Service Matcher", "/tools/service-matcher", "Interactive quiz to match participants with recommended services"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 6. ADMIN PANEL - USER MANUAL
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Admin Panel (CRM) \u2013 User Manual", level=1)
    add_divider(doc)

    add_styled_paragraph(
        doc,
        "The Admin Panel is a comprehensive CRM (Customer Relationship Management) system accessible at "
        "/admin. It provides tools for managing leads, referrals, blog content, gallery media, analytics, and user settings.",
        font_size=11, space_after=8
    )

    add_info_box(doc,
                 "How to Access the Admin Panel",
                 "1. Navigate to: jschoicegroup.com.au/admin/login\n"
                 "2. Enter your email and password\n"
                 "3. You will be redirected to the Dashboard\n"
                 "4. Use the sidebar menu to navigate between modules")

    # 6.1 Dashboard
    add_heading_styled(doc, "6.1 Dashboard Overview", level=2)
    add_styled_paragraph(
        doc,
        "The Dashboard is the first page you see after logging in. It provides a quick snapshot of your "
        "platform's activity and key metrics.",
        font_size=11, space_after=8
    )

    add_styled_paragraph(doc, "Dashboard Components:", font_size=11, bold=True, space_after=4)
    add_bullet_list(doc, [
        "Total Leads \u2013 Count of all incoming inquiries from all sources",
        "Total Referrals \u2013 Count of referral-specific submissions",
        "Blog Posts \u2013 Number of published blog articles",
        "Conversion Rate \u2013 Percentage of leads that moved to 'converted' status",
        "Recent Activity Feed \u2013 Live timeline of latest actions (new leads, status changes, emails)",
        "Engagement Chart \u2013 Visual representation of lead activity over time",
    ])

    # 6.2 Lead Management
    add_heading_styled(doc, "6.2 Lead Management", level=2)
    add_styled_paragraph(
        doc,
        "The Lead Management module is the core of your CRM. Every form submission from the website "
        "(contact form, service matcher, budget calculator, NDIS tools) automatically creates a lead here.",
        font_size=11, space_after=8
    )

    add_styled_paragraph(doc, "Lead Pipeline Statuses:", font_size=11, bold=True, space_after=4)
    add_styled_table(doc,
        ["Status", "Color", "Description"],
        [
            ["New", "Blue", "Just received \u2013 needs first contact"],
            ["Contacted", "Yellow", "Initial outreach has been made"],
            ["Qualified", "Purple", "Confirmed as genuine NDIS participant"],
            ["Converted", "Green", "Successfully became a client"],
            ["Lost", "Red", "Did not proceed or went elsewhere"],
        ],
        col_widths=[1.2, 0.8, 4.5]
    )

    add_styled_paragraph(doc, "How to Manage Leads:", font_size=11, bold=True, space_after=4)
    add_bullet_list(doc, [
        "View All Leads: Navigate to \u2018Leads\u2019 in the sidebar. Use search bar to find specific leads.",
        "Filter Leads: Filter by status (New, Contacted, Qualified, etc.) or priority (Low, Normal, High, Urgent).",
        "View Lead Detail: Click on any lead to open detailed view with contact info, service interests, and activity log.",
        "Update Status: In the lead detail view, change the status dropdown to move leads through the pipeline.",
        "Add Notes: Use the activity section to add internal notes, log calls, or record meeting outcomes.",
        "Send Email: Click the email icon to compose and send an email directly to the lead from the CRM.",
        "Create Tasks: Assign follow-up tasks with due dates, priority, and type (call, email, meeting).",
        "Export Data: Use the \u2018Export\u2019 button to download all leads as CSV or Excel file.",
        "Delete Lead: Use the delete option (with confirmation) to permanently remove a lead.",
    ])

    add_styled_paragraph(doc, "Lead Sources (Automatically Tracked):", font_size=11, bold=True, space_after=4)
    add_styled_table(doc,
        ["Source", "Origin"],
        [
            ["contact_form", "Main contact page form submission"],
            ["referral", "Referral form submission"],
            ["service_matcher", "Interactive service matching quiz"],
            ["budget_calculator", "NDIS budget calculator"],
            ["blog", "Blog page inquiry"],
            ["ndis_tools", "NDIS price guide / tools usage"],
            ["phone", "Manual entry from phone call"],
            ["walk_in", "Manual entry from walk-in visit"],
        ],
        col_widths=[2.0, 4.5]
    )

    # 6.3 Referral Management
    add_heading_styled(doc, "6.3 Referral Management", level=2)
    add_styled_paragraph(
        doc,
        "The Referral module works identically to Lead Management but is filtered to show only "
        "referral submissions. Referrals contain additional details such as participant conditions, "
        "service hours requested, cultural considerations, guardian info, and NDIS plan type.",
        font_size=11, space_after=8
    )

    add_styled_paragraph(doc, "Referral Form Fields Captured:", font_size=11, bold=True, space_after=4)
    add_bullet_list(doc, [
        "Referrer type (self or someone else)",
        "Participant details (name, DOB, gender, address, phone, email, NDIS number)",
        "Guardian/nominee information",
        "Cultural background (country of birth, interpreter needs, cultural/religious considerations)",
        "Aboriginal or Torres Strait Islander identification",
        "Primary & secondary service requests with weekly hours",
        "Disability conditions and additional notes",
        "Consultation type preferences (in-clinic, in-home, telehealth, community)",
        "Contact person for appointments",
        "NDIS plan type (NDIA managed, plan managed, self-managed)",
    ])

    # 6.4 Blog Management
    add_heading_styled(doc, "6.4 Blog Management", level=2)
    add_styled_paragraph(
        doc,
        "Create and manage blog content to engage your audience and improve SEO.",
        font_size=11, space_after=8
    )

    add_styled_paragraph(doc, "How to Create a Blog Post:", font_size=11, bold=True, space_after=4)
    add_bullet_list(doc, [
        "Navigate to Blog in the sidebar and click \u2018New Post\u2019",
        "Enter the title \u2013 a URL slug will be automatically generated",
        "Write your content using the rich text editor (supports bold, italic, headings, links, images)",
        "Upload a featured image using the image upload area",
        "Add tags (comma-separated) and select a category",
        "Fill in SEO fields: Meta Title, Meta Description, Canonical URL",
        "Choose status: Draft (save for later), Published (go live now), or Scheduled (set a future date)",
        "Click \u2018Publish\u2019 or \u2018Save Draft\u2019",
    ])

    add_styled_paragraph(doc, "Blog Post Statuses:", font_size=11, bold=True, space_after=4)
    add_styled_table(doc,
        ["Status", "Description"],
        [
            ["Draft", "Saved but not visible to the public"],
            ["Published", "Live and visible on the website"],
            ["Scheduled", "Will auto-publish on the set date and time"],
            ["Archived", "Hidden from public, preserved for reference"],
        ],
        col_widths=[1.5, 5.0]
    )

    # 6.5 Gallery Management
    add_heading_styled(doc, "6.5 Gallery Management", level=2)
    add_styled_paragraph(
        doc,
        "Manage photo gallery content that showcases your services and team.",
        font_size=11, space_after=8
    )

    add_styled_paragraph(doc, "How to Manage Gallery:", font_size=11, bold=True, space_after=4)
    add_bullet_list(doc, [
        "Navigate to Gallery in the sidebar",
        "Click \u2018Add New\u2019 to create a gallery card",
        "Upload up to 5 images per gallery card",
        "Add a title, description, and assign a category",
        "Set the display order number to control positioning",
        "Save the gallery item \u2013 it will appear on the public gallery page",
        "Edit or delete existing items using the action buttons",
    ])

    # 6.6 Analytics
    add_heading_styled(doc, "6.6 Analytics & Reports", level=2)
    add_styled_paragraph(
        doc,
        "The Analytics module provides visual insights into your lead pipeline and business performance.",
        font_size=11, space_after=8
    )

    add_styled_paragraph(doc, "Available Reports:", font_size=11, bold=True, space_after=4)
    add_bullet_list(doc, [
        "Lead Growth Over Time \u2013 Area chart showing new leads per day/week/month",
        "Source Breakdown \u2013 Pie chart showing where leads come from (contact form, referral, tools, etc.)",
        "Status Distribution \u2013 Bar chart showing leads by pipeline status",
        "Conversion Rate \u2013 Percentage of leads converted to clients",
        "Time Range Filters \u2013 Filter data by 7 days, 30 days, 90 days, or all time",
    ])

    # 6.7 Settings
    add_heading_styled(doc, "6.7 Settings & Profile", level=2)
    add_styled_paragraph(doc, "Manage your admin account and preferences.", font_size=11, space_after=8)
    add_bullet_list(doc, [
        "Update display name and job title",
        "Upload/change profile avatar",
        "Update phone number and email",
        "Toggle email notification preferences (new lead alerts, daily digests, etc.)",
        "Change password",
        "View account role and permissions",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 7. BACKEND MODULES & API
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "7. Backend Modules & API", level=1)
    add_divider(doc)

    add_styled_paragraph(
        doc,
        "The backend is built using Next.js API routes (serverless functions). Below is a complete inventory "
        "of all API modules and their endpoints.",
        font_size=11, space_after=10
    )

    # 7.1 Auth
    add_heading_styled(doc, "7.1 Authentication Module", level=2)
    add_styled_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["POST", "/api/auth/logout", "End the current admin session"],
            ["GET", "/api/auth/session", "Retrieve current authenticated session"],
        ],
        col_widths=[0.8, 2.5, 3.2]
    )

    # 7.2 Lead/CRM
    add_heading_styled(doc, "7.2 Lead / CRM Module", level=2)
    add_styled_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/api/leads", "List all leads (with filters, search, pagination)"],
            ["POST", "/api/leads", "Create a new lead from any form"],
            ["GET", "/api/leads/[id]", "Get single lead details"],
            ["PATCH", "/api/leads/[id]", "Update lead status, priority, notes"],
            ["DELETE", "/api/leads/[id]", "Delete a lead permanently"],
            ["GET", "/api/leads/[id]/activities", "Get activity log for a lead"],
            ["POST", "/api/leads/[id]/activities", "Add note, log call, or status change"],
            ["GET", "/api/leads/[id]/tasks", "Get tasks assigned to a lead"],
            ["POST", "/api/leads/[id]/tasks", "Create follow-up task for a lead"],
            ["POST", "/api/leads/[id]/email", "Send email to a lead"],
            ["GET", "/api/leads/export", "Export all leads as CSV/Excel"],
            ["PATCH", "/api/tasks/[taskId]", "Update or complete a task"],
            ["DELETE", "/api/tasks/[taskId]", "Delete a task"],
        ],
        col_widths=[0.8, 2.5, 3.2]
    )

    # 7.3 Blog
    add_heading_styled(doc, "7.3 Blog Module", level=2)
    add_styled_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/api/blog", "List blog posts with filtering"],
            ["POST", "/api/blog", "Create a new blog post"],
            ["GET", "/api/blog/[slug]", "Get single blog post by slug"],
            ["PATCH", "/api/blog/[slug]", "Update an existing blog post"],
            ["DELETE", "/api/blog/[slug]", "Delete a blog post"],
            ["POST", "/api/blog/upload", "Upload blog images to storage"],
            ["GET", "/api/blog/categories", "Get all blog categories"],
            ["POST", "/api/blog/scheduler", "Auto-publish scheduled posts"],
        ],
        col_widths=[0.8, 2.5, 3.2]
    )

    # 7.4 Gallery
    add_heading_styled(doc, "7.4 Gallery Module", level=2)
    add_styled_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/api/gallery", "List all gallery items"],
            ["POST", "/api/gallery", "Create a new gallery card"],
            ["GET", "/api/gallery/[id]", "Get a single gallery item"],
            ["PATCH", "/api/gallery/[id]", "Update gallery item details"],
            ["DELETE", "/api/gallery/[id]", "Delete a gallery item"],
            ["POST", "/api/gallery/upload", "Upload gallery images to storage"],
        ],
        col_widths=[0.8, 2.5, 3.2]
    )

    # 7.5 NDIS Tools
    add_heading_styled(doc, "7.5 NDIS Tools Module", level=2)
    add_styled_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/api/ndis/search", "Full-text search across 5,000+ NDIS items"],
            ["GET", "/api/ndis/autocomplete", "Real-time autocomplete for search bar"],
            ["GET", "/api/ndis/categories", "Get all 15 NDIS support categories"],
            ["GET", "/api/ndis/item/[code]", "Get detailed NDIS item with all region prices"],
            ["GET", "/api/ndis/services", "Get JS Choice service offerings"],
            ["POST", "/api/ndis/leads", "Create lead from NDIS tools interaction"],
        ],
        col_widths=[0.8, 2.5, 3.2]
    )

    # 7.6 Email
    add_heading_styled(doc, "7.6 Email Module (Resend)", level=2)
    add_styled_paragraph(
        doc,
        "The email system uses Resend API to send branded HTML emails for:",
        font_size=11, space_after=6
    )
    add_bullet_list(doc, [
        "Admin Notifications \u2013 Instant alert when a new lead/referral is submitted",
        "Client Confirmation \u2013 Auto-reply to the person who submitted a form",
        "Direct Email \u2013 Send emails to leads directly from the CRM communication panel",
        "Follow-Up Reminders \u2013 Email reminders for overdue tasks",
    ])

    # 7.7 Analytics
    add_heading_styled(doc, "7.7 Analytics Module", level=2)
    add_styled_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/api/analytics/overview", "Dashboard stats (totals, rates, recent activity)"],
            ["GET", "/api/analytics/leads", "Lead charts data (growth, sources, status)"],
        ],
        col_widths=[0.8, 2.5, 3.2]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 8. DATABASE ARCHITECTURE
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "8. Database Architecture", level=1)
    add_divider(doc)

    add_styled_paragraph(
        doc,
        "The platform uses Supabase (PostgreSQL) as its primary database. Below is the complete schema "
        "with all 10 tables and their purposes.",
        font_size=11, space_after=10
    )

    add_styled_table(doc,
        ["Table", "Records", "Description"],
        [
            ["ndis_support_items", "5,000+", "Complete NDIS price list with region-specific pricing for all Australian states"],
            ["ndis_categories", "15", "NDIS support categories (Core, Capital, Capacity Building)"],
            ["jschoice_services", "13", "JS Choice service offerings mapped to NDIS categories"],
            ["leads", "Dynamic", "All leads/contacts from forms, tools, phone, walk-ins with full CRM data"],
            ["lead_activities", "Dynamic", "Activity log: status changes, notes, emails, calls, meetings per lead"],
            ["lead_tasks", "Dynamic", "Follow-up tasks with due dates, priority, type, completion status"],
            ["blog_posts", "Dynamic", "Blog articles with rich content, SEO fields, scheduling, categories"],
            ["blog_categories", "Dynamic", "Blog category taxonomy with parent support"],
            ["gallery_items", "Dynamic", "Gallery cards with up to 5 images, categories, display order"],
            ["profiles", "Dynamic", "Admin user profiles with roles, avatar, notification preferences"],
        ],
        col_widths=[1.8, 0.8, 3.9]
    )

    add_heading_styled(doc, "Row Level Security (RLS)", level=2)
    add_styled_paragraph(
        doc,
        "All tables have Row Level Security enabled for data protection:",
        font_size=11, space_after=6
    )
    add_bullet_list(doc, [
        "Public users can READ: NDIS items, categories, services, published blog posts, gallery items",
        "Public users can INSERT: leads (form submissions) \u2013 but cannot read other leads",
        "Authenticated admins can CRUD: All tables through the admin panel",
        "Service role (API): Full access for server-side operations",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 9. FEATURES SUMMARY
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "9. Features Summary", level=1)
    add_divider(doc)

    features = [
        ["CRM & Lead Management", "Full pipeline tracking, activity logging, task management, email communication, export"],
        ["Multi-Step Referral Form", "6-section referral with participant details, cultural info, service requests, NDIS plan details"],
        ["NDIS Price Guide", "Search 5,000+ support items, filter by category, view region-specific pricing"],
        ["NDIS Budget Calculator", "Add items to budget, set frequency, calculate annual costs, PDF export"],
        ["Service Matcher Quiz", "Interactive questionnaire that recommends services based on participant needs"],
        ["Blog System", "Rich text editor, image uploads, scheduling, SEO fields, categories, tags"],
        ["Photo Gallery", "Multi-image cards, categories, display ordering, Supabase storage"],
        ["Contact Forms", "Contact page form, referral form \u2013 all auto-create CRM leads"],
        ["Email Automation", "Admin alerts, client confirmations, direct emails from CRM, HTML templates"],
        ["Analytics Dashboard", "Lead growth charts, source breakdown, status distribution, conversion rates"],
        ["User Management", "Admin profiles, avatars, roles (admin/manager/staff), notification preferences"],
        ["SEO Optimization", "Dynamic meta tags, canonical URLs, 16 location pages, blog categories"],
        ["Responsive Design", "Fully responsive across mobile, tablet, and desktop"],
        ["Animations & UX", "Framer Motion animations, glass-morphism, parallax effects, smooth transitions"],
        ["PDF Generation", "Budget calculator exports to professional PDF with branding"],
        ["Data Export", "Lead export to CSV/Excel for external analysis"],
    ]

    add_styled_table(doc,
        ["Feature", "Description"],
        features,
        col_widths=[2.2, 4.3]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 10. SECURITY & COMPLIANCE
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "10. Security & Compliance", level=1)
    add_divider(doc)

    add_bullet_list(doc, [
        "Authentication: Supabase Auth with secure session management and cookie-based tokens",
        "Row Level Security (RLS): Database-level access control on all tables",
        "Role-Based Access: Admin, Manager, Staff roles with appropriate permissions",
        "HTTPS: All traffic encrypted with TLS/SSL via Vercel edge network",
        "Input Validation: Server-side validation on all API endpoints",
        "CSRF Protection: Next.js built-in CSRF protections",
        "Environment Variables: Sensitive keys stored securely, never exposed to client",
        "Image Upload Security: File type and size validation on all uploads",
        "Rate Limiting: API protection against abuse (via Vercel)",
        "NDIS Compliance: Platform designed to support NDIS Quality & Safeguards requirements",
        "Privacy: Participant data handling aligned with Australian Privacy Principles",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 11. DEPLOYMENT & INFRASTRUCTURE
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "11. Deployment & Infrastructure", level=1)
    add_divider(doc)

    add_styled_table(doc,
        ["Component", "Service", "Details"],
        [
            ["Hosting", "Vercel", "Auto-scaling serverless platform with global CDN edge network"],
            ["Database", "Supabase", "Managed PostgreSQL with real-time, auth, and storage"],
            ["Image Storage", "Supabase Storage", "S3-compatible object storage for uploads"],
            ["Email Delivery", "Resend", "Transactional email API with high deliverability"],
            ["Domain", "jschoicegroup.com.au", "Custom domain with SSL certificate"],
            ["CI/CD", "Vercel Git Integration", "Auto-deploy on push to main branch"],
            ["Monitoring", "Vercel Analytics", "Performance monitoring and error tracking"],
        ],
        col_widths=[1.3, 1.8, 3.4]
    )

    add_heading_styled(doc, "Environment Variables", level=2)
    add_styled_table(doc,
        ["Variable", "Purpose"],
        [
            ["NEXT_PUBLIC_SUPABASE_URL", "Supabase project URL (public)"],
            ["NEXT_PUBLIC_SUPABASE_ANON_KEY", "Supabase anonymous key (public)"],
            ["SUPABASE_SERVICE_ROLE_KEY", "Supabase service role key (server-only, secret)"],
            ["RESEND_API_KEY", "Resend email service API key (server-only, secret)"],
            ["NEXT_PUBLIC_APP_URL", "Application base URL"],
        ],
        col_widths=[3.0, 3.5]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 12. SUPPORT & CONTACT
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "12. Support & Contact", level=1)
    add_divider(doc)

    add_styled_paragraph(
        doc,
        "For technical support, platform issues, or feature requests regarding the JS Choice Group "
        "digital platform, please contact:",
        font_size=11, space_after=12
    )

    add_info_box(doc,
                 "JS Choice Group Pty Ltd",
                 "Phone: 0421 622 262\n"
                 "Website: jschoicegroup.com.au\n"
                 "Location: Point Cook, Melbourne, VIC\n"
                 "ABN: Registered NDIS Provider")

    for _ in range(3):
        doc.add_paragraph()

    add_divider(doc, PRIMARY_HEX)

    add_styled_paragraph(doc, "END OF DOCUMENT",
                          font_size=14, bold=True, color=DARK,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=20, space_after=8)

    add_styled_paragraph(doc, "JS Choice Group Pty Ltd \u2013 User Manual & Final Report v1.2.2",
                          font_size=10, color=GRAY,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    add_styled_paragraph(doc, "Prepared: February 16, 2026 | Confidential",
                          font_size=9, color=GRAY, italic=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_divider(doc, SECONDARY_HEX)

    # ─── SAVE ──────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(__file__), "JS_CHOICE_USER_MANUAL.docx")
    doc.save(output_path)
    print(f"\nDocument generated successfully!")
    print(f"Saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_document()
