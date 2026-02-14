"""
Convert Markdown to DOCX with proper formatting
Handles headings, tables, code blocks, lists, and text formatting
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def parse_inline_formatting(text, paragraph):
    """Parse and apply inline formatting (bold, italic, code, links)."""
    # Pattern for: **bold**, *italic*, `code`, [text](url)
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[([^\]]+)\]\(([^)]+)\)|[^*`\[]+)'

    parts = re.finditer(pattern, text)

    for match in parts:
        content = match.group(0)

        if content.startswith('**') and content.endswith('**'):
            # Bold
            run = paragraph.add_run(content[2:-2])
            run.bold = True
        elif content.startswith('*') and content.endswith('*'):
            # Italic
            run = paragraph.add_run(content[1:-1])
            run.italic = True
        elif content.startswith('`') and content.endswith('`'):
            # Inline code
            run = paragraph.add_run(content[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(220, 20, 60)
        elif content.startswith('[') and '](' in content:
            # Link
            link_text = match.group(2)
            link_url = match.group(3)
            if link_url and link_text:
                try:
                    add_hyperlink(paragraph, link_text, link_url)
                except:
                    paragraph.add_run(f"{link_text} ({link_url})")
        else:
            # Plain text
            if content.strip():
                paragraph.add_run(content)

def convert_markdown_to_docx(md_file, docx_file):
    """Convert Markdown file to DOCX with formatting."""

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle code blocks
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_lines = []
                i += 1
                continue
            else:
                # End of code block
                in_code_block = False
                if code_block_lines:
                    # Add code block to document
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run('\n'.join(code_block_lines))
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(9)
                    code_run.font.color.rgb = RGBColor(40, 40, 40)
                    # Add gray background
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'F5F5F5')
                    code_para._element.get_or_add_pPr().append(shading_elm)
                code_block_lines = []
                i += 1
                continue

        if in_code_block:
            code_block_lines.append(line.rstrip())
            i += 1
            continue

        # Handle tables
        if stripped.startswith('|') and '|' in stripped:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            # Check if next line is not a table line
            if i < len(lines) and not lines[i].strip().startswith('|'):
                # Process table
                process_table(doc, table_lines)
                in_table = False
                table_lines = []
            continue

        # Handle headings
        if stripped.startswith('#'):
            level = len(re.match(r'^#+', stripped).group())
            text = stripped[level:].strip()

            heading_para = doc.add_heading(text, level=min(level, 9))

            # Custom formatting for heading
            if level == 1:
                heading_para.runs[0].font.size = Pt(24)
                heading_para.runs[0].font.color.rgb = RGBColor(0, 0, 139)
            elif level == 2:
                heading_para.runs[0].font.size = Pt(20)
                heading_para.runs[0].font.color.rgb = RGBColor(25, 25, 112)
            elif level == 3:
                heading_para.runs[0].font.size = Pt(16)
                heading_para.runs[0].font.color.rgb = RGBColor(70, 130, 180)

        # Handle horizontal rules
        elif stripped == '---' or stripped == '***':
            doc.add_paragraph('_' * 50)

        # Handle unordered lists
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            para = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(text, para)

        # Handle ordered lists
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            para = doc.add_paragraph(style='List Number')
            parse_inline_formatting(text, para)

        # Handle blockquotes
        elif stripped.startswith('>'):
            text = stripped[1:].strip()
            para = doc.add_paragraph()
            parse_inline_formatting(text, para)
            para.paragraph_format.left_indent = Inches(0.5)
            # Add gray background
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E8E8E8')
            para._element.get_or_add_pPr().append(shading_elm)

        # Handle empty lines
        elif not stripped:
            if i > 0 and lines[i-1].strip():  # Only add space after content
                doc.add_paragraph()

        # Handle regular paragraphs
        else:
            if stripped:
                para = doc.add_paragraph()
                parse_inline_formatting(stripped, para)

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

def process_table(doc, table_lines):
    """Process and add a table to the document."""
    # Filter out separator lines
    data_lines = [line for line in table_lines if not re.match(r'^\|[\s\-:|]+\|$', line)]

    if len(data_lines) < 2:
        return

    # Parse table data
    rows = []
    for line in data_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last
        rows.append(cells)

    if not rows:
        return

    # Create table
    num_cols = len(rows[0])
    num_rows = len(rows)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Fill table
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]

            # Parse inline formatting in cell
            para = cell.paragraphs[0]
            para.text = ''  # Clear default text
            parse_inline_formatting(cell_text, para)

            # Make header row bold
            if i == 0:
                for run in para.runs:
                    run.bold = True
                # Add shading to header
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '4F81BD')
                cell._element.get_or_add_tcPr().append(shading_elm)
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)

if __name__ == '__main__':
    md_file = 'SYSTEM_DESIGN_DOCUMENT_COMPLETE.md'
    docx_file = 'SYSTEM_DESIGN_DOCUMENT_COMPLETE.docx'

    print(f"Converting {md_file} to {docx_file}...")
    print("This may take a minute for large documents...")

    convert_markdown_to_docx(md_file, docx_file)

    print(f"\nConversion complete!")
    print(f"Output file: {docx_file}")
